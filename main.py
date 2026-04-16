from contextlib import asynccontextmanager
from datetime import datetime
import os
import json
import asyncio
from dotenv import load_dotenv
from docx import Document
from google import genai
from typing import List

from fastapi import FastAPI, HTTPException, Depends, Request, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, Field
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime

# Initialize environment & LLM Client
load_dotenv()
try:
    gemini_client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
except Exception:
    gemini_client = None
from sqlalchemy.orm import declarative_base, sessionmaker, Session

# ---------------------------------------------------------
# Database Setup Setup
# ---------------------------------------------------------
SQLALCHEMY_DATABASE_URL = "sqlite:///./gemini_sources.db"
engine = create_engine(
    SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False}
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class GeminiSource(Base):
    __tablename__ = "gemini_sources"

    id = Column(Integer, primary_key=True, index=True)
    title = Column(String, index=True)
    content = Column(Text)
    timestamp = Column(DateTime, default=datetime.utcnow)
    source_url = Column(String, index=True)

class GeneratedReport(Base):
    __tablename__ = "generated_reports"

    id = Column(Integer, primary_key=True, index=True)
    report_data = Column(Text) # Stored serialized JSON
    timestamp = Column(DateTime, default=datetime.utcnow)

# ---------------------------------------------------------
# Pydantic Models for Validation
# ---------------------------------------------------------
class SourceCreate(BaseModel):
    title: str
    content: str
    source_url: str

class PipelineStep(BaseModel):
    prompt_text: str = Field(description="The exact text to copy-paste into Antigravity/Cursor/Claude.")
    why: str = Field(description="Explanation of why this step exists.")
    expectation: str = Field(description="What the project state should look like after running it.")
    error_warnings: str = Field(description="What red flags to look out for regarding errors.")

class AnalysisSchema(BaseModel):
    summary: str = Field(description="Brief product overview.")
    swot: str = Field(description="Strengths, Weaknesses, Opportunities, Threats.")
    market_analysis: str = Field(description="Top competitors, service differences, and Blue Ocean viability.")
    cost_benefit: str = Field(description="Financial and operational tradeoff of building it.")
    viability_score: int = Field(description="Integer from 0-100 indicating sure-fire success vs flop.")
    vibe_coding_pipeline: list[PipelineStep] = Field(description="Sequential timeline of implementation prompts.")

class AnalyzeRequest(BaseModel):
    target_platform: str = "Antigravity"

class BulkDeleteRequest(BaseModel):
    source_ids: list[int]

class OutlineSchema(BaseModel):
    chapters: list[str]

class SourceResponse(BaseModel):
    id: int
    title: str
    content: str
    timestamp: datetime
    source_url: str

    class Config:
        from_attributes = True

# ---------------------------------------------------------
# FastAPI App Setup
# ---------------------------------------------------------
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Create the database tables on startup
    Base.metadata.create_all(bind=engine)
    yield
    # Any cleanup could go here

app = FastAPI(title="Gemini Sources API", lifespan=lifespan)

# Enable CORS for extensions
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust this in production to specific origins (e.g., your extension's ID)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files and set up templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# ---------------------------------------------------------
# WebSocket Manager
# ---------------------------------------------------------
class ConnectionManager:
    def __init__(self):
        self.active_connections: list[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)

    async def broadcast(self, message: str):
        for connection in self.active_connections:
            try:
                await connection.send_text(message)
            except Exception:
                pass

manager = ConnectionManager()

# Database Dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---------------------------------------------------------
# Endpoints
# ---------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
def read_root(request: Request):
    """Serve the NotebookLM-style frontend interaction page."""
    return templates.TemplateResponse("index.html", {"request": request})

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(websocket)

@app.get("/api/sources")
def get_sources(db: Session = Depends(get_db)):
    """Return all ingested sources ordered by most recent."""
    sources = db.query(GeminiSource).order_by(GeminiSource.timestamp.desc()).all()
    return sources

@app.delete("/api/sources/{source_id}")
async def delete_source(source_id: int, db: Session = Depends(get_db)):
    db_source = db.query(GeminiSource).filter(GeminiSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Source not found")
    db.delete(db_source)
    db.commit()
    await manager.broadcast("new_source")
    return {"status": "success"}

@app.post("/api/sources/bulk-delete")
async def bulk_delete_sources(req: BulkDeleteRequest, db: Session = Depends(get_db)):
    if not req.source_ids:
        return {"status": "success", "deleted_count": 0}
        
    db.query(GeminiSource).filter(GeminiSource.id.in_(req.source_ids)).delete(synchronize_session=False)
    db.commit()
    await manager.broadcast("new_source")
    return {"status": "success", "deleted_count": len(req.source_ids)}

@app.post("/ingest", response_model=SourceResponse)
async def ingest_source(source: SourceCreate, db: Session = Depends(get_db)):
    """Accepts JSON and saves it to the DB."""
    db_source = GeminiSource(
        title=source.title,
        content=source.content,
        source_url=source.source_url,
        timestamp=datetime.utcnow()
    )
    db.add(db_source)
    db.commit()
    db.refresh(db_source)
    
    # Notify connected real-time UI components
    await manager.broadcast("new_source")
    
    return db_source

@app.post("/api/analyze")
async def analyze_sources(req: AnalyzeRequest, db: Session = Depends(get_db)):
    if not gemini_client:
        raise HTTPException(status_code=500, detail="Gemini SDK improperly configured. Check API key.")
        
    sources = db.query(GeminiSource).order_by(GeminiSource.timestamp.asc()).all()
    if not sources:
        raise HTTPException(status_code=400, detail="No sources found to analyze.")
        
    context_text = "\n\n".join([f"Source: {s.title}\n{s.content}" for s in sources])
    
    prompt = f"""
    You are an elite Software Architect and Business Analyst.
    Analyze the following brainstorm notes and output a rigorous structured analysis based on the exact JSON schema requested.
    The target vibe coding platform the user will use is [{req.target_platform}]. 
    Please tailor the 'vibe_coding_pipeline' prompts specifically for this platform so they can copy paste them directly into the tool.
    
    Brainstorm Context:
    {context_text}
    """
    
    try:
        response = gemini_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config={
                'response_mime_type': 'application/json',
                'response_schema': AnalysisSchema,
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
        
    generated_json = response.text
    
    # Save generic generated report
    new_report = GeneratedReport(report_data=generated_json, timestamp=datetime.utcnow())
    db.add(new_report)
    db.commit()
    
    # Notify UI
    await manager.broadcast("new_report")
    
    return json.loads(generated_json)
    
@app.get("/api/reports/latest")
def get_latest_report(db: Session = Depends(get_db)):
    report = db.query(GeneratedReport).order_by(GeneratedReport.timestamp.desc()).first()
    if report:
        return json.loads(report.report_data)
    return None

async def generate_architect_report(source_texts: str, platform: str):
    await manager.broadcast(json.dumps({"type": "progress", "message": "Initializing Architect Framework...", "progress": 10}))
    
    outline_prompt = f"""
    You are an expert technical and business architect. Analyze these raw notes and generate a comprehensive structural outline for a massive professional blueprint.
    Target platform: {platform}
    
    Raw Notes:
    {source_texts}
    """
    
    try:
        outline_res = gemini_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=outline_prompt,
            config={
                'response_mime_type': 'application/json',
                'response_schema': OutlineSchema,
            },
        )
        outline_data = json.loads(outline_res.text)
        chapters = outline_data.get('chapters', [])
    except Exception as e:
        await manager.broadcast(json.dumps({"type": "error", "message": f"Outline generation failed: {str(e)}"}))
        return

    doc = Document()
    doc.add_heading('SynapseIP Master Blueprint', 0)
    
    total_chapters = len(chapters)
    await manager.broadcast(json.dumps({"type": "progress", "message": f"Outline verified. Writing {total_chapters} chapters...", "progress": 20}))
    
    for i, chapter_title in enumerate(chapters):
        prog = 20 + int((i / total_chapters) * 70)
        await manager.broadcast(json.dumps({"type": "progress", "message": f"Drafting Chapter {i+1}: {chapter_title}...", "progress": prog}))
        
        chapter_prompt = f"""
        You are an expert architect writing a deep-dive, professional chapter for a massive blueprint document.
        
        Document Outline: {json.dumps(chapters)}
        Current Chapter to Write: '{chapter_title}'
        Target Platform: {platform}
        
        Based ONLY on the following raw notes, write a highly detailed, 4-page equivalent professional business and technical guide for this specific chapter. 
        Use professional maturity. DO NOT write an introduction to the whole document, just write the chapter itself.
        
        Raw Notes:
        {source_texts}
        """
        
        try:
            chap_res = gemini_client.models.generate_content(
                model='gemini-2.5-flash',
                contents=chapter_prompt
            )
            doc.add_heading(chapter_title, level=1)
            doc.add_paragraph(chap_res.text)
            doc.add_page_break()
        except Exception as e:
            print(f"Skipping chapter {chapter_title} due to error: {e}")
        
        # THROTTLE FOR 429
        await asyncio.sleep(4)

    os.makedirs('static/reports', exist_ok=True)
    file_path = "static/reports/SynapseIP_Master_Plan.docx"
    doc.save(file_path)
    
    await manager.broadcast(json.dumps({
        "type": "architect_complete",
        "message": "Master Document compiled and saved.",
        "progress": 100,
        "download_url": f"/{file_path}"
    }))

@app.post("/api/architect/start")
async def start_architect(req: AnalyzeRequest, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    sources = db.query(GeminiSource).all()
    if not sources:
        raise HTTPException(status_code=400, detail="No sources available. Sync some datanodes first.")
        
    combined_text = "\n\n---\n\n".join([f"TITLE: {s.title}\n{s.content}" for s in sources])
    background_tasks.add_task(generate_architect_report, combined_text, req.target_platform)
    
    return {"status": "started", "message": "Architect pipeline initiated."}

# To run: uvicorn main:app --reload
